import json
import pandas as pd
from io import BytesIO
from docx import Document
from fpdf import FPDF
from docx.shared import RGBColor, Inches
from datetime import datetime
import zipfile
import xml.etree.ElementTree as ET
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import math
import os
from io import BytesIO as _BytesIO
from typing import Tuple, List, Dict, Optional

def parse_agent_output(output_text: str):
    """Try to parse agent output as JSON -> DataFrame"""
    try:
        if isinstance(output_text, (dict, list)):
            data = output_text
        else:
            data = json.loads(output_text)
        if isinstance(data, list):
            return pd.DataFrame(data)
        elif isinstance(data, dict):
            return pd.DataFrame([data])
        else:
            return pd.DataFrame()
    except Exception:
        return pd.DataFrame([{"Raw Output": str(output_text)}])


# ----------------- Basic exports -----------------

def export_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8") if not df.empty else b""

def export_excel(df: pd.DataFrame) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Traceability")
    return output.getvalue()

def export_word(df: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading("Traceability Export", level=1)
    if df.empty:
        doc.add_paragraph("No data available")
    else:
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def export_pdf(df: pd.DataFrame) -> bytes:
    def sanitize_for_pdf(text: str) -> str:
        """Replace Unicode characters that can't be encoded in latin-1 with ASCII equivalents."""
        replacements = {
            '\u2264': '<=',  # ≤
            '\u2265': '>=',  # ≥
            '\u2260': '!=',  # ≠
            '\u00b1': '+/-', # ±
            '\u2192': '->',  # →
            '\u2190': '<-',  # ←
            '\u2794': '->',  # ➔
            '\u27a1': '->',  # ➡
            '\u279c': '->',  # ➜
            '\u279e': '->',  # ➞
            '\u2964': '->',  # ⥤
            '\u21d2': '=>',  # ⇒
            '\u21d4': '<=>',  # ⇔
            '\u00b0': 'deg', # °
            '\u03bc': 'u',   # µ (micro)
            '\u2022': '*',   # •
            '\u2013': '-',   # –
            '\u2014': '--',  # —
            '\u2018': "'",   # '
            '\u2019': "'",   # '
            '\u201c': '"',   # "
            '\u201d': '"',   # "
        }
        for uni, ascii_rep in replacements.items():
            text = text.replace(uni, ascii_rep)
        # Fallback: encode to latin-1 with error replacement
        try:
            text.encode('latin-1')
            return text
        except UnicodeEncodeError:
            return text.encode('latin-1', errors='replace').decode('latin-1')
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    if df.empty:
        pdf.cell(0, 10, "No data available", ln=True)
    else:
        col_width = pdf.w / (len(df.columns) + 1)
        row_height = pdf.font_size * 1.5
        for col in df.columns:
            pdf.cell(col_width, row_height, sanitize_for_pdf(str(col)), border=1)
        pdf.ln(row_height)
        for _, row in df.iterrows():
            for item in row:
                pdf.cell(col_width, row_height, sanitize_for_pdf(str(item)), border=1)
            pdf.ln(row_height)
    raw = pdf.output(dest='S')  # pyfpdf returns a str (latin-1) or bytes depending on version
    if isinstance(raw, str):
        return raw.encode('latin-1')
    return raw

# ----------------- Styled RTM Exports -----------------

def export_excel_styled(df: pd.DataFrame, sheet_name="RTM") -> bytes:
    """Export DataFrame to Excel with wrapped text & adaptive column widths.

    Strategy:
    - Write headers + data
    - Wrap text in every cell (so long requirements / test steps flow vertically)
    - Compute width heuristic from the longest line (split on whitespace) but clamp to a max
    - Keep header row centered
    """
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4F81BD")

    # Header
    ws.append(list(df.columns))
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Rows (coerce complex types to string for openpyxl)
    def _cell(v):
        if v is None:
            return ""
        # stringify lists/dicts/sets and other non-primitive values
        if isinstance(v, (list, tuple, set, dict)):
            try:
                return json.dumps(v, ensure_ascii=False)
            except Exception:
                return str(v)
        return v if isinstance(v, (int, float, bool)) else str(v)
    for row in df.itertuples(index=False):
        ws.append([_cell(v) for v in row])

    # Apply wrap + width heuristic
    MAX_WIDTH = 60  # excel units (~chars)
    MIN_WIDTH = 10
    for col in ws.columns:
        col_letter = col[0].column_letter
        longest = 0
        for cell in col:
            val = cell.value
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if val is None:
                continue
            text = str(val)
            # Consider longest word to avoid excessive width
            parts = text.split()
            if parts:
                lw = max(len(p) for p in parts)
                if lw > longest:
                    longest = lw
            # fallback to line length average
            avg = len(text) / max(1, text.count(" ") + 1)
            longest = max(longest, int(avg))
        ws.column_dimensions[col_letter].width = max(MIN_WIDTH, min(MAX_WIDTH, longest + 2))

    output = BytesIO()
    wb.save(output)
    return output.getvalue()

def save_excel_styled_to_path(df: pd.DataFrame, filepath: str, sheet_name: str = "Data") -> str:
    """Save DataFrame to an .xlsx file on disk with readable formatting.

    - Creates parent folders if they don't exist
    - Wrap text for all cells
    - Bold header row with background fill
    - Freeze header row
    - Adaptive column widths (heuristic)
    Returns the absolute file path saved.
    """
    if df is None:
        df = pd.DataFrame()
    # Ensure directory exists
    dirpath = os.path.dirname(os.path.abspath(filepath))
    if dirpath and not os.path.exists(dirpath):
        os.makedirs(dirpath, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4F81BD")

    # Header
    ws.append(list(df.columns))
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Rows (coerce complex types to string for openpyxl)
    def _cell(v):
        if v is None:
            return ""
        if isinstance(v, (list, tuple, set, dict)):
            try:
                return json.dumps(v, ensure_ascii=False)
            except Exception:
                return str(v)
        return v if isinstance(v, (int, float, bool)) else str(v)
    for row in df.itertuples(index=False):
        ws.append([_cell(v) for v in row])

    # Freeze header row
    ws.freeze_panes = "A2"

    # Apply wrap + width heuristic
    MAX_WIDTH = 60
    MIN_WIDTH = 10
    for col in ws.columns:
        col_letter = col[0].column_letter
        longest = 0
        for cell in col:
            val = cell.value
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if val is None:
                continue
            text = str(val)
            parts = text.split()
            if parts:
                lw = max(len(p) for p in parts)
                if lw > longest:
                    longest = lw
            avg = len(text) / max(1, text.count(" ") + 1)
            longest = max(longest, int(avg))
        ws.column_dimensions[col_letter].width = max(MIN_WIDTH, min(MAX_WIDTH, longest + 2))

    abs_path = os.path.abspath(filepath)
    tmp_path = abs_path + ".tmp"
    # Primary write to temp then atomic replace
    wb.save(tmp_path)
    try:
        os.replace(tmp_path, abs_path)
    except Exception:
        # Best-effort fallback: direct save
        wb.save(abs_path)
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass

    # Verify on disk, if not, fallback to bytes-based writer
    try:
        ok = os.path.exists(abs_path) and os.path.getsize(abs_path) > 0
    except Exception:
        ok = False
    if not ok:
        try:
            data = export_excel_styled(df, sheet_name)
            with open(abs_path, 'wb') as f:
                f.write(data)
        except Exception:
            pass
    return abs_path

from docx.shared import Pt
def export_word_styled(df: pd.DataFrame, project="WHALE Project", version="1.0", author="AI Co-pilot", logo_path=None) -> bytes:
    doc = Document()
    # Cover
    if logo_path:
        try:
            doc.add_picture(logo_path, width=Inches(2))
        except Exception:
            pass
    doc.add_heading("Requirements Traceability Matrix (RTM)", 0)
    doc.add_paragraph(f"Project: {project}")
    doc.add_paragraph(f"Version: {version}")
    doc.add_paragraph(f"Author: {author}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_page_break()

    # Table
    doc.add_heading("RTM", level=1)
    if df.empty:
        doc.add_paragraph("No data available")
    else:
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col
            if hdr_cells[i].paragraphs and hdr_cells[i].paragraphs[0].runs:
                hdr_cells[i].paragraphs[0].runs[0].bold = True
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                txt = str(value)
                # Replace explicit newlines with paragraph breaks for readability
                paragraphs = txt.split('\n')
                first = True
                for p in paragraphs:
                    if first:
                        row_cells[i].text = p
                        first = False
                    else:
                        row_cells[i].add_paragraph(p)
                for para in row_cells[i].paragraphs:
                    para.paragraph_format.space_after = Pt(2)
                    for run in para.runs:
                        run.font.size = Pt(10)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def export_pdf_styled(df: pd.DataFrame, project="WHALE Project", version="1.0", author="AI Co-pilot", logo_path=None) -> bytes:
    """PDF export with wrapped table cells and adaptive column widths.

    Approach:
    - Compute target widths from header & content character lengths
    - Scale to available width
    - Wrap text for each cell using multi_cell; align rows vertically
    """
    
    # Suppress FPDF warnings about invalid float values in PDF objects
    import warnings
    warnings.filterwarnings('ignore', message='.*invalid float value.*')
    
    def sanitize_for_pdf(text: str) -> str:
        """Replace Unicode characters that can't be encoded in latin-1 with ASCII equivalents."""
        replacements = {
            '\u2264': '<=',  # ≤
            '\u2265': '>=',  # ≥
            '\u2260': '!=',  # ≠
            '\u00b1': '+/-', # ±
            '\u2192': '->',  # →
            '\u2190': '<-',  # ←
            '\u2794': '->',  # ➔
            '\u27a1': '->',  # ➡
            '\u279c': '->',  # ➜
            '\u279e': '->',  # ➞
            '\u2964': '->',  # ⥤
            '\u21d2': '=>',  # ⇒
            '\u21d4': '<=>',  # ⇔
            '\u00b0': 'deg', # °
            '\u03bc': 'u',   # µ (micro)
            '\u2022': '*',   # •
            '\u2013': '-',   # –
            '\u2014': '--',  # —
            '\u2018': "'",   # '
            '\u2019': "'",   # '
            '\u201c': '"',   # "
            '\u201d': '"',   # "
        }
        for uni, ascii_rep in replacements.items():
            text = text.replace(uni, ascii_rep)
        # Fallback: encode to latin-1 with error replacement
        try:
            text.encode('latin-1')
            return text
        except UnicodeEncodeError:
            return text.encode('latin-1', errors='replace').decode('latin-1')
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if logo_path:
        try:
            pdf.image(logo_path, x=80, y=20, w=50)
            pdf.ln(60)
        except Exception:
            pass
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, sanitize_for_pdf("Requirements Traceability Matrix (RTM)"), ln=True, align="C")
    pdf.ln(10)
    pdf.set_font("Arial", "", 11)
    meta = [f"Project: {sanitize_for_pdf(project)}", f"Version: {sanitize_for_pdf(version)}", f"Author: {sanitize_for_pdf(author)}", f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}"]
    for m in meta:
        pdf.cell(0, 8, sanitize_for_pdf(m), ln=True)
    pdf.ln(4)

    if not df.empty:
        # Column width calculation
        pdf.set_font("Arial", "", 8)
        char_w = pdf.get_string_width("M")  # rough monospace approximation
        col_lengths = []
        for c in df.columns:
            max_len = len(str(c))
            for v in df[c].astype(str):
                max_len = max(max_len, min(120, len(v)))  # clamp for huge text
            col_lengths.append(max_len)
        total_chars = sum(col_lengths)
        usable_w = pdf.w - 2 * pdf.l_margin
        # Raw proportional widths
        col_widths = [max(15, usable_w * (cl / total_chars)) for cl in col_lengths]
        # Ensure sum fits exactly (rounding adjustment)
        scale = usable_w / sum(col_widths)
        col_widths = [w * scale for w in col_widths]

        # Header
        pdf.set_font("Arial", "B", 8)
        pdf.set_fill_color(79,129,189)
        pdf.set_text_color(255,255,255)
        h_line = 6
        for w, col in zip(col_widths, df.columns):
            x_before = pdf.get_x()
            pdf.multi_cell(w, h_line, sanitize_for_pdf(str(col)), border=1, align='C', fill=True)
            pdf.set_xy(x_before + w, pdf.get_y() - h_line)  # keep on same row
        pdf.ln(h_line)
        pdf.set_text_color(0,0,0)
        pdf.set_font("Arial", "", 7)

        line_height = 4
        for _, row in df.iterrows():
            # Precompute wrapped lines per cell
            wrapped_cells = []
            max_lines = 1
            for text, w in zip(row.astype(str), col_widths):
                # Sanitize text before processing
                text = sanitize_for_pdf(text)
                words = text.split()
                lines = []
                current = ""
                for word in words:
                    trial = (current + " " + word).strip()
                    if pdf.get_string_width(trial) <= (w - 1):
                        current = trial
                    else:
                        if current:
                            lines.append(current)
                        current = word
                if current:
                    lines.append(current)
                if not lines:
                    lines = [""]
                wrapped_cells.append(lines)
                max_lines = max(max_lines, len(lines))
            row_height = line_height * max_lines
            y_start = pdf.get_y()
            x_start = pdf.get_x()
            for lines, w in zip(wrapped_cells, col_widths):
                x_cell = pdf.get_x()
                y_cell = pdf.get_y()
                # Draw border rectangle manually to align multi-line cell
                pdf.rect(x_cell, y_cell, w, row_height)
                for i, ln in enumerate(lines):
                    pdf.set_xy(x_cell + 1, y_cell + i * line_height)
                    pdf.cell(w - 2, line_height, ln, align='L')
                pdf.set_xy(x_cell + w, y_cell)
            pdf.set_xy(x_start, y_start + row_height)

    raw = pdf.output(dest='S')
    if isinstance(raw, str):
        return raw.encode('latin-1')
    return raw

# ----------------- REQIF Export -----------------

def export_reqif(df: pd.DataFrame, spec_name="WHALE Requirements") -> bytes:
    reqif = ET.Element("REQ-IF")
    header = ET.SubElement(reqif, "THE-HEADER")
    tool_id = ET.SubElement(header, "TOOL-ID")
    tool_id.text = "WHALE-AI"
    created = ET.SubElement(header, "CREATION-TIME")
    created.text = datetime.now().isoformat()
    core_content = ET.SubElement(reqif, "CORE-CONTENT")
    spec_objects = ET.SubElement(core_content, "SPEC-OBJECTS")
    for i, row in df.iterrows():
        spec_object = ET.SubElement(spec_objects, "SPEC-OBJECT", {"IDENTIFIER": f"REQ-{i+1}"})
        values = ET.SubElement(spec_object, "VALUES")
        for col, val in row.items():
            attr_val = ET.SubElement(values, "ATTRIBUTE-VALUE", {"ATTRIBUTE-DEFINITION-REF": col})
            val_elem = ET.SubElement(attr_val, "THE-VALUE")
            val_elem.text = str(val)
    specifications = ET.SubElement(core_content, "SPECIFICATIONS")
    spec = ET.SubElement(specifications, "SPECIFICATION", {"IDENTIFIER": "SPEC-1"})
    long_name = ET.SubElement(spec, "LONG-NAME")
    long_name.text = spec_name
    xml_str = ET.tostring(reqif, encoding="utf-8", method="xml")
    return xml_str

# ----------------- ZIP bundle -----------------

def export_all_as_zip(tables: dict) -> bytes:
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zf:
        for name, df in tables.items():
            if df is None or df.empty:
                continue
            # CSV
            zf.writestr(f"{name}.csv", df.to_csv(index=False))
            # Excel
            zf.writestr(f"{name}.xlsx", export_excel(df))
            # Word
            zf.writestr(f"{name}.docx", export_word(df))
            # PDF
            zf.writestr(f"{name}.pdf", export_pdf(df))
            # REQIF
            zf.writestr(f"{name}.reqif", export_reqif(df, spec_name=name))
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

# ----------------- RTM builder -----------------

def build_traceability_matrix(cust_sys1, sys1_sys2, review, sys2_sys5):
    rtm = cust_sys1.copy() if not cust_sys1.empty else pd.DataFrame()
    if not sys1_sys2.empty:
        # try merging on SYS.1 Req. ID
        key = "SYS.1 Req. ID" if "SYS.1 Req. ID" in cust_sys1.columns and "SYS.1 Req. ID" in sys1_sys2.columns else None
        if key:
            rtm = cust_sys1.merge(sys1_sys2, on=key, how="left")
        else:
            rtm = pd.concat([cust_sys1, sys1_sys2], axis=1)
    if not review.empty and "SYS.2 Req. ID" in rtm.columns:
        rtm = rtm.merge(review, on="SYS.2 Req. ID", how="left")
    if not sys2_sys5.empty and "SYS.2 Req. ID" in rtm.columns:
        rtm = rtm.merge(sys2_sys5, on="SYS.2 Req. ID", how="left")
    return rtm


# ----------------- File Import Utilities (DOCX/PDF/XLSX/REQIF) -----------------

def read_docx_text_from_bytes(b: bytes) -> str:
    try:
        from docx import Document as _Doc
        f = _BytesIO(b)
        doc = _Doc(f)
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(paras)
    except Exception:
        return ""

def read_pdf_text_from_bytes(b: bytes) -> str:
    try:
        # Requires PyPDF2
        from PyPDF2 import PdfReader
        reader = PdfReader(_BytesIO(b))
        texts = []
        for page in getattr(reader, 'pages', []):
            try:
                t = page.extract_text() or ""
                if t.strip():
                    texts.append(t)
            except Exception:
                continue
        return "\n\n".join(texts)
    except Exception:
        return ""

def read_xlsx_df_from_bytes(b: bytes):
    try:
        return pd.read_excel(_BytesIO(b))
    except Exception:
        return pd.DataFrame()

def read_reqif_df_from_bytes(b: bytes) -> pd.DataFrame:
    """Parse REQIF XML to DataFrame. Tries multiple REQIF formats and XML structures."""
    try:
        xml = ET.parse(_BytesIO(b))
        root = xml.getroot()
        items: List[Dict[str, str]] = []
        
        # Try standard REQIF structure: SPEC-OBJECTS -> SPEC-OBJECT -> VALUES -> ATTRIBUTE-VALUE
        spec_objects = root.findall('.//SPEC-OBJECT')
        if not spec_objects:
            # Try without namespace or different paths
            spec_objects = root.findall('.//{*}SPEC-OBJECT')
        
        for so in spec_objects:
            row: Dict[str, str] = {}
            # Try to get identifier
            obj_id = so.attrib.get('IDENTIFIER') or so.attrib.get('ID') or ''
            if obj_id:
                row['ID'] = obj_id
            
            # Look for attribute values
            attr_vals = so.findall('.//ATTRIBUTE-VALUE')
            if not attr_vals:
                attr_vals = so.findall('.//{*}ATTRIBUTE-VALUE')
            
            for av in attr_vals:
                key = (av.attrib.get('ATTRIBUTE-DEFINITION-REF') or 
                       av.attrib.get('DEFINITION-REF') or 
                       av.attrib.get('THE-ATTRIBUTE') or 
                       'Field')
                
                # Try multiple value element names
                val_elem = (av.find('THE-VALUE') or 
                           av.find('{*}THE-VALUE') or 
                           av.find('VALUE'))
                
                if val_elem is not None:
                    val = val_elem.text or ''
                else:
                    # Sometimes value is directly in the element
                    val = av.text or ''
                
                if key and val:
                    row[key] = val
            
            if row:
                items.append(row)
        
        # If no structured data found, try to extract any text content as fallback
        if not items:
            all_text = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    all_text.append(elem.text.strip())
            if all_text:
                # Create a single-column dataframe with all text
                items = [{"Requirement": txt} for txt in all_text if len(txt) > 10]
        
        return pd.DataFrame(items) if items else pd.DataFrame()
    except Exception as e:
        # Return empty DataFrame on any parsing error
        return pd.DataFrame()

def _df_to_sys1_text(df: pd.DataFrame) -> str:
    if df is None or df.empty:
        return ""
    parts = []
    # Try to honor common columns
    id_col = next((c for c in df.columns if 'customer' in c.lower() and 'id' in c.lower()), None)
    req_col = next((c for c in df.columns if 'requirement' in c.lower() and 'sys.1' not in c.lower() and 'sys1' not in c.lower()), None)
    for _, row in df.iterrows():
        ident = str(row.get(id_col, '')).strip() if id_col else ''
        body = str(row.get(req_col, '')).strip() if req_col else ''
        if ident or body:
            prefix = f"{ident}: " if ident else ''
            parts.append(prefix + body if body else prefix)
        else:
            # fallback: join non-empty values
            vals = [str(v).strip() for v in row.astype(str).tolist() if str(v).strip()]
            parts.append(" | ".join(vals))
    return "\n\n".join([p for p in parts if p])

def _df_to_records(df: pd.DataFrame) -> List[Dict[str, str]]:
    if df is None or df.empty:
        return []
    return df.fillna("").astype(str).to_dict(orient='records')

def _norm_header_key(s: str) -> str:
    # Lowercase, remove non-alphanumerics; keep dots for SYS.1/SYS.2 then remove them
    import re as _re
    s = s or ""
    s = s.replace("\n", " ")
    s = _re.sub(r"\s+", " ", s).strip().lower()
    s = s.replace(".", "").replace("_", " ")
    return _re.sub(r"[^a-z0-9]+", "", s)

def _canonical_map_for_agent(agent_key: str) -> Dict[str, str]:
    # Map normalized header -> canonical field name
    m: Dict[str, str] = {}
    def add(names: List[str], canon: str):
        for n in names:
            m[_norm_header_key(n)] = canon
    # Common
    add(["domain"], "Domain")
    add(["priority"], "Priority")
    add(["rationale"], "Rationale")
    add(["requirement status","status"], "Requirement Status")
    # Agent 1 (SYS.1)
    add(["customer req id","customer req_id","customer requirement id","customer id"], "Customer Req. ID")
    add(["customer requirement"], "Customer Requirement")
    add(["sys1 req id","sys 1 req id","sys.1 req id","sys.1 req _id","sys1 requirement id"], "SYS.1 Req. ID")
    add(["sys1 requirement","sys.1 requirement"], "SYS.1 Requirement")
    # Agent 2 (SYS.2)
    add(["sys2 req id","sys 2 req id","sys.2 req id","sys.2 req _id"], "SYS.2 Req. ID")
    add(["sys2 requirement","sys.2 requirement"], "SYS.2 Requirement")
    # Preserve backward compatibility: existing data may have been normalized to "TYPE".
    # Map common variants to "Type" (preferred) while also mapping legacy normalized key "type" to both.
    add(["type","requirement type","req type","functional or nonfunctional"], "Type")
    add(["verification method","system qualification test (sys5)","method"], "Verification Method")
    add(["verification criteria","criteria"], "Verification Criteria")
    # Test artifacts (SYS5)
    add(["test case id","tc id","test_id"], "Test Case ID")
    add(["description","test description"], "Description")
    add(["preconditions","precondition"], "Preconditions")
    add(["test steps","steps","test_steps"], "Test Steps")
    add(["expected result","expected","expected_result"], "Expected Result")
    add(["pass/fail criteria","pass fail criteria","passfail criteria"], "Pass/Fail Criteria")
    add(["test level","level"], "Test Level")
    add(["safety goal link","safety goal","sg link","sg"], "Safety Goal Link")
    # Agent 3 (Review) (kept below for ordering clarity)
    add(["review feedback","feedback"], "Review Feedback")
    add(["compliance check","compliance"], "Compliance Check")
    add(["suggested improvement","improvement","correction/enhancement"], "Suggested Improvement")
    add(["smart check"], "SMART Check")
    return m

def normalize_df_for_agent(agent_key: str, df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    cmap = _canonical_map_for_agent(agent_key)
    new_cols = []
    for c in df.columns:
        canon = cmap.get(_norm_header_key(str(c)), c)
        new_cols.append(canon)
    df = df.copy()
    df.columns = new_cols
    return df

def normalize_records_for_agent(agent_key: str, records: List[Dict[str, str]]) -> List[Dict[str, str]]:
    if not records:
        return records
    cmap = _canonical_map_for_agent(agent_key)
    out: List[Dict[str, str]] = []
    for r in records:
        newr: Dict[str, str] = {}
        for k, v in r.items():
            canon = cmap.get(_norm_header_key(str(k)), k)
            newr[canon] = v
        out.append(newr)
    return out

def load_file_for_agent(file_bytes: bytes, filename: str, agent_key: str) -> Tuple[Optional[str], Optional[List[Dict[str, str]]]]:
    """Given uploaded file bytes + name, return (text, records) best suited for agent.

    - For SYS1: prefer text (concatenated requirements)
    - For SYS2/Review/SYS5: prefer records (list of dicts) if possible; else return text
    """
    name = filename.lower()
    text = ""
    records: Optional[List[Dict[str, str]]] = None
    if name.endswith('.txt') or name.endswith('.json'):
        try:
            raw = file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            raw = ""
        if name.endswith('.json'):
            try:
                data = json.loads(raw)
                if isinstance(data, list):
                    records = [x if isinstance(x, dict) else {"Raw": str(x)} for x in data]
                elif isinstance(data, dict):
                    # unwrap wrappers like {"SYS1": [...]} or {"SYS2": [...]}
                    for k in ("SYS1","SYS2","Review","SYS5","data"):
                        if k in data and isinstance(data[k], list):
                            recs = data[k]
                            records = [x if isinstance(x, dict) else {"Raw": str(x)} for x in recs]
                            break
                else:
                    text = raw
            except Exception:
                text = raw
        else:
            text = raw
    elif name.endswith('.docx'):
        text = read_docx_text_from_bytes(file_bytes)
    elif name.endswith('.pdf'):
        text = read_pdf_text_from_bytes(file_bytes)
    elif name.endswith('.xlsx'):
        df = read_xlsx_df_from_bytes(file_bytes)
        df = normalize_df_for_agent(agent_key, df)
        if agent_key == 'SYS1':
            text = _df_to_sys1_text(df)
        else:
            records = _df_to_records(df)
    elif name.endswith('.reqif'):
        df = read_reqif_df_from_bytes(file_bytes)
        df = normalize_df_for_agent(agent_key, df)
        if agent_key == 'SYS1':
            text = _df_to_sys1_text(df)
        else:
            records = _df_to_records(df)
    else:
        # Unknown extension; best-effort text
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            text = ""

    # Normalize outputs per agent intention
    if agent_key == 'SYS1':
        return (text, None)
    # For record-oriented agents, normalize keys if we have records
    if records is not None:
        records = normalize_records_for_agent(agent_key, records)
    return (text if not records else None, records)
